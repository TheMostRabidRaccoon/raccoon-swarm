"""Read-only Google Drive observation surface for the swarm.

The swarm previously received only a filename manifest for Drive. That proves an
artifact exists but does not make its information cognitively available.

This module uses an already-configured rclone Google Drive remote to expose:

- drive_status  — observation-surface/config status, never credentials;
- drive_search  — Google Drive full-text query via `rclone backend query`;
- drive_read    — fetch one matched file by Drive ID to a temporary LOCAL path,
                  extract readable text, then delete the temporary copy.

No upload, delete, move, rename, permission, or remote-write operation exists in
this surface. Configure the rclone remote itself with Drive's read-only scope when
possible. A missing remote/binary is an unavailable observation route, not a claim
about participant capability.
"""
from __future__ import annotations

import html
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET


DEFAULT_TIMEOUT = 30
DEFAULT_MAX_RESULTS = 10
DEFAULT_MAX_CHARS = 16_000
MAX_DOWNLOAD_BYTES = 25 * 1024 * 1024
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_ARCHIVE_MEMBER_BYTES = 50 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 2_000_000

_STOPWORDS = {
    "a", "an", "and", "are", "as", "at", "be", "but", "by", "for", "from",
    "had", "has", "have", "how", "i", "in", "is", "it", "me", "my", "of",
    "on", "or", "our", "that", "the", "their", "them", "this", "to", "was",
    "we", "were", "what", "when", "where", "which", "who", "why", "with", "you",
    "your",
}
_DRIVE_ID_RE = re.compile(r"^[A-Za-z0-9_-]{8,200}$")


def _remote() -> str:
    return (os.getenv("RRI_DRIVE_REMOTE") or "").strip()


def _timeout() -> int:
    try:
        return max(5, min(int(os.getenv("RRI_DRIVE_TIMEOUT", DEFAULT_TIMEOUT)), 120))
    except (TypeError, ValueError):
        return DEFAULT_TIMEOUT


def _run(args: list[str], timeout: int | None = None) -> subprocess.CompletedProcess:
    """Run rclone without shell interpolation. Kept as a seam for unit tests."""
    return subprocess.run(
        args,
        capture_output=True,
        text=True,
        timeout=timeout or _timeout(),
        check=False,
    )


def _configured() -> tuple[bool, str]:
    remote = _remote()
    if not remote:
        return False, "RRI_DRIVE_REMOTE is not configured"
    if shutil.which("rclone") is None:
        return False, "rclone is not installed on this runtime surface"
    return True, "configured"


def status(probe: bool = False) -> dict:
    configured, reason = _configured()
    out = {
        "ok": configured,
        "surface": "google-drive/read-only-observation",
        "configured": configured,
        "reason": reason,
        "remote": _remote() if configured else None,
        "remote_write_actuator": "not exposed on this surface",
        "recommended_rclone_scope": "drive.readonly",
        "max_transfer_bytes": MAX_DOWNLOAD_BYTES,
    }
    if not configured or not probe:
        return out

    proc = _run(["rclone", "backend", "query", _remote(), "trashed = false"])
    out["probe_ok"] = proc.returncode == 0
    if proc.returncode != 0:
        out["probe_error"] = (proc.stderr or proc.stdout or "rclone query failed")[:500]
    return out


def _escape_drive_literal(value: str) -> str:
    return (value or "").replace("\\", "\\\\").replace("'", "\\'")


def _terms(query: str, limit: int = 7) -> list[str]:
    seen: set[str] = set()
    terms: list[str] = []
    pattern = r"[A-Za-z0-9]+(?:['’][A-Za-z0-9]+)*(?:[_.-][A-Za-z0-9]+)*"
    for raw_token in re.findall(pattern, query or ""):
        token = raw_token.replace("’", "'")
        low = token.lower().strip("._-")
        if len(low) < 3 or low in _STOPWORDS or low in seen:
            continue
        seen.add(low)
        terms.append(token)
        if len(terms) >= limit:
            break
    return terms


def _drive_query(query: str) -> str:
    terms = _terms(query)
    if not terms:
        escaped = _escape_drive_literal((query or "").strip())
        return f"trashed = false and fullText contains '{escaped}'"
    ors = " or ".join(
        f"fullText contains '{_escape_drive_literal(term)}'" for term in terms
    )
    return f"trashed = false and ({ors})"


def search(query: str, max_results: int = DEFAULT_MAX_RESULTS) -> dict:
    configured, reason = _configured()
    if not configured:
        return {"ok": False, "query": query, "error": reason, "results": []}
    if not (query or "").strip():
        return {"ok": False, "query": query, "error": "query is empty", "results": []}

    max_results = max(1, min(int(max_results or DEFAULT_MAX_RESULTS), 30))
    drive_q = _drive_query(query)
    proc = _run(["rclone", "backend", "query", _remote(), drive_q])
    if proc.returncode != 0:
        return {
            "ok": False,
            "query": query,
            "drive_query": drive_q,
            "error": (proc.stderr or proc.stdout or "rclone query failed")[:1000],
            "results": [],
        }
    try:
        rows = json.loads(proc.stdout or "[]")
    except json.JSONDecodeError as exc:
        return {"ok": False, "query": query, "error": f"invalid rclone JSON: {exc}", "results": []}
    if not isinstance(rows, list):
        rows = []

    q_terms = {t.lower() for t in _terms(query, limit=12)}
    q_lower = (query or "").lower().strip()

    def rank(row: dict) -> tuple[int, str]:
        name = str(row.get("name") or "").lower()
        exact = 5 if q_lower and q_lower in name else 0
        token_hits = sum(1 for t in q_terms if t in name)
        return exact + token_hits, str(row.get("modifiedTime") or "")

    rows.sort(key=rank, reverse=True)
    results = []
    for row in rows[:max_results]:
        if not isinstance(row, dict):
            continue
        results.append({
            "id": row.get("id"),
            "name": row.get("name"),
            "mime_type": row.get("mimeType"),
            "modified": row.get("modifiedTime"),
            "created": row.get("createdTime"),
            "size": row.get("size"),
            "web_view_link": row.get("webViewLink"),
        })
    return {
        "ok": True,
        "query": query,
        "drive_query": drive_q,
        "results": results,
        "total_returned": len(results),
        "surface": "google-drive/read-only-observation",
    }


def _assert_archive_safe(path: Path) -> None:
    """Bound Office zip-container expansion before a parser opens it."""
    try:
        with zipfile.ZipFile(path) as zf:
            infos = zf.infolist()
            total = sum(max(0, info.file_size) for info in infos)
            largest = max((max(0, info.file_size) for info in infos), default=0)
    except zipfile.BadZipFile as exc:
        raise ValueError("invalid Office archive") from exc
    if total > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
        raise ValueError(
            f"archive expansion exceeds limit ({MAX_ARCHIVE_UNCOMPRESSED_BYTES} bytes)"
        )
    if largest > MAX_ARCHIVE_MEMBER_BYTES:
        raise ValueError(
            f"archive member exceeds limit ({MAX_ARCHIVE_MEMBER_BYTES} bytes)"
        )


def _text_from_docx(path: Path) -> str:
    from docx import Document

    _assert_archive_safe(path)
    doc = Document(str(path))
    parts: list[str] = []
    chars = 0
    for paragraph in doc.paragraphs:
        if not paragraph.text:
            continue
        parts.append(paragraph.text)
        chars += len(paragraph.text)
        if chars >= MAX_EXTRACTED_TEXT_CHARS:
            break
    if chars < MAX_EXTRACTED_TEXT_CHARS:
        for table in doc.tables:
            for row in table.rows:
                line = "\t".join(cell.text for cell in row.cells)
                parts.append(line)
                chars += len(line)
                if chars >= MAX_EXTRACTED_TEXT_CHARS:
                    break
            if chars >= MAX_EXTRACTED_TEXT_CHARS:
                break
    return "\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS]


def _text_from_pdf(path: Path) -> str:
    import fitz

    doc = fitz.open(str(path))
    parts: list[str] = []
    chars = 0
    try:
        for page in doc:
            text = page.get_text("text")
            remaining = MAX_EXTRACTED_TEXT_CHARS - chars
            if remaining <= 0:
                break
            parts.append(text[:remaining])
            chars += min(len(text), remaining)
    finally:
        doc.close()
    return "\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS]


def _text_from_xlsx(path: Path, max_rows_per_sheet: int = 500) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[XLSX text extraction unavailable: install openpyxl]"

    _assert_archive_safe(path)
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    chars = 0
    try:
        for ws in wb.worksheets:
            header = f"# Sheet: {ws.title}"
            parts.append(header)
            chars += len(header)
            for i, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if i > max_rows_per_sheet:
                    marker = f"[truncated after {max_rows_per_sheet} rows]"
                    parts.append(marker)
                    chars += len(marker)
                    break
                vals = ["" if v is None else str(v) for v in row]
                if any(vals):
                    line = "\t".join(vals)
                    remaining = MAX_EXTRACTED_TEXT_CHARS - chars
                    if remaining <= 0:
                        break
                    parts.append(line[:remaining])
                    chars += min(len(line), remaining)
            if chars >= MAX_EXTRACTED_TEXT_CHARS:
                break
    finally:
        wb.close()
    return "\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS]


def _text_from_pptx(path: Path) -> str:
    _assert_archive_safe(path)
    parts: list[str] = []
    chars = 0
    try:
        with zipfile.ZipFile(path) as zf:
            names = sorted(
                n for n in zf.namelist()
                if re.match(r"ppt/slides/slide\d+\.xml$", n)
            )
            for name in names:
                root = ET.fromstring(zf.read(name))
                texts = [
                    html.unescape(el.text or "")
                    for el in root.iter()
                    if el.tag.endswith("}t")
                ]
                if texts:
                    line = " ".join(texts)
                    remaining = MAX_EXTRACTED_TEXT_CHARS - chars
                    if remaining <= 0:
                        break
                    parts.append(line[:remaining])
                    chars += min(len(line), remaining)
    except (zipfile.BadZipFile, ET.ParseError, OSError):
        return ""
    return "\n".join(parts)[:MAX_EXTRACTED_TEXT_CHARS]


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".txt", ".md", ".json", ".csv", ".tsv", ".log", ".py", ".yaml", ".yml", ".html", ".htm"}:
        with open(path, "r", encoding="utf-8", errors="replace") as handle:
            return handle.read(MAX_EXTRACTED_TEXT_CHARS)
    if ext == ".docx":
        return _text_from_docx(path)
    if ext == ".pdf":
        return _text_from_pdf(path)
    if ext == ".xlsx":
        return _text_from_xlsx(path)
    if ext == ".pptx":
        return _text_from_pptx(path)
    return ""


def read(file_id: str, max_chars: int = DEFAULT_MAX_CHARS) -> dict:
    """Fetch one Drive file by ID into a bounded temp surface and extract text."""
    configured, reason = _configured()
    if not configured:
        return {"ok": False, "file_id": file_id, "error": reason}
    if not _DRIVE_ID_RE.match(file_id or ""):
        return {"ok": False, "file_id": file_id, "error": "invalid Drive file id"}
    max_chars = max(500, min(int(max_chars or DEFAULT_MAX_CHARS), 60_000))

    with tempfile.TemporaryDirectory(prefix="rri-drive-read-") as td:
        dest = Path(td)
        limit = f"{MAX_DOWNLOAD_BYTES}B"
        # Enforce the byte ceiling DURING transfer, not after the file has already
        # occupied local disk. The post-transfer stat check remains defense in depth.
        proc = _run([
            "rclone", "backend", "copyid", _remote(), file_id, f"{dest.as_posix()}/",
            "--drive-export-formats", "txt,csv,pdf",
            "--max-size", limit,
            "--max-transfer", limit,
            "--cutoff-mode", "HARD",
        ], timeout=max(_timeout(), 45))
        if proc.returncode != 0:
            return {
                "ok": False,
                "file_id": file_id,
                "error": (proc.stderr or proc.stdout or "rclone copyid failed")[:1000],
            }
        files = [p for p in dest.rglob("*") if p.is_file()]
        if not files:
            return {"ok": False, "file_id": file_id, "error": "Drive fetch returned no local file"}
        path = max(files, key=lambda p: p.stat().st_size)
        size = path.stat().st_size
        if size > MAX_DOWNLOAD_BYTES:
            return {
                "ok": False,
                "file_id": file_id,
                "name": path.name,
                "bytes": size,
                "error": f"file exceeds read-surface limit ({MAX_DOWNLOAD_BYTES} bytes)",
            }
        try:
            text = _extract_text(path)
        except Exception as exc:
            return {
                "ok": False,
                "file_id": file_id,
                "name": path.name,
                "bytes": size,
                "error": f"text extraction failed: {type(exc).__name__}: {exc}",
            }
        if not text:
            return {
                "ok": True,
                "file_id": file_id,
                "name": path.name,
                "bytes": size,
                "content": "",
                "text_available": False,
                "note": "file was fetched read-only but no supported text representation was extracted",
            }
        truncated = len(text) > max_chars
        return {
            "ok": True,
            "file_id": file_id,
            "name": path.name,
            "bytes": size,
            "content": text[:max_chars],
            "chars": len(text),
            "truncated": truncated,
            "text_available": True,
            "surface": "google-drive/read-only-observation",
        }


def tool_definitions() -> dict:
    return {
        "drive_status": {
            "description": (
                "Inspect the swarm's read-only Google Drive observation surface. Reports whether "
                "an rclone Drive remote is configured; never returns credentials. A missing remote "
                "means this observation route is not configured, not that Drive reasoning is beyond "
                "the participant's capability."
            ),
            "input_schema": {
                "type": "object",
                "properties": {"probe": {"type": "boolean", "description": "Run a live metadata-only reachability probe. Default false."}},
            },
            "dispatch": status,
        },
        "drive_search": {
            "description": (
                "Search Google Drive's provider index, including indexed file text and metadata, "
                "through the configured read-only observation surface. Returns Drive file IDs and "
                "metadata; use drive_read on a relevant result to obtain its actual content."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Natural-language/content query."},
                    "max_results": {"type": "integer", "description": "Maximum results, 1-30. Default 10."},
                },
                "required": ["query"],
            },
            "dispatch": search,
        },
        "drive_read": {
            "description": (
                "Fetch one Google Drive result by its Drive file ID to a bounded temporary LOCAL "
                "read surface, extract text from supported document formats, and delete the "
                "temporary copy afterward. Transfer and compressed-document expansion limits are "
                "enforced before content is admitted to the observation surface. This tool exposes "
                "no remote write/delete/move operation."
            ),
            "input_schema": {
                "type": "object",
                "properties": {
                    "file_id": {"type": "string", "description": "Drive file ID returned by drive_search."},
                    "max_chars": {"type": "integer", "description": "Maximum extracted characters returned, 500-60000. Default 16000."},
                },
                "required": ["file_id"],
            },
            "dispatch": read,
        },
    }
