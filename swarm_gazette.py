"""The Daily Burrow + Play Gazette — newspaper-style Conductor emails.

Two products, one honesty rule:

- **Daily Burrow** — a once-a-day mechanical digest of the last 24h: session
  ledger, flags/blockers, email-channel health (closer gap counts surfaced
  verbatim), Joy runs, receipts. Fired by scripts/run_daily_burrow.py on the
  swarm-burrow.timer.
- **Play Gazette** — a front-page edition for PLAY-shaped sessions, with an
  optional DOCX attachment. Fired by the closer immediately after a play
  session's receipts land, and swept by the daily runner as a backstop.

Honesty invariants (learned the hard way across this repo):

- Everything reported is read from receipts on disk — scorecard-*.json,
  corpus/corpus-*.json, closer-digest-*.md, joy/runs/**, logs/emails.log.
  No model self-report is trusted or re-narrated.
- A closer gap>0 is repeated verbatim: N emails were owed and did NOT send.
  The Burrow exists partly to make those visible (the inbox subject lines
  were carrying the gap counts long before anyone decoded them).
- "attached" is only ever claimed for a file that existed and was hashed
  BEFORE the send (swarm_mail.send_operational refuses otherwise) — the
  attachment version of "announcing a write is not a write".
- Unmeasured is never reported as fine: a missing scorecard field renders
  as "?" in the ledger, not 0.

Stdlib + swarm_filestore + swarm_gate (both stdlib-only). python-docx is
optional — no DOCX means the email still goes out, honestly unattached.
"""
from __future__ import annotations

import json
import logging
import os
import re
from datetime import datetime, timedelta
from pathlib import Path

import swarm_filestore
import swarm_gate

logger = logging.getLogger("SwarmVault")

GAZETTE_VERSION = 1

# Filestore lanes for the persisted markdown editions. The persisted edition is
# also the idempotency marker: a play gazette exists iff its file exists.
DAILY_LANE = "artifacts/gazettes/daily"
PLAY_LANE = "artifacts/gazettes/play"

# PLAY detection is mechanical: the declared purpose tag wins; otherwise a
# short, documented marker list over the query text. Extend by PR, not vibes.
PLAY_MARKERS = (
    "play session",
    "pure play",
    "freeplay",
    "free-play",
    "zero deliverables",
    "woodland council",
    "the bouncer",
    "chitterverse",
    "joy mode",
)


def enabled() -> bool:
    return os.getenv("RRI_GAZETTE_ENABLED", "true").strip().lower() not in ("0", "false", "no", "off")


def is_play_session(query: str) -> bool:
    """Mechanical PLAY classifier: [SESSION_PURPOSE: creative-production] or a
    known play marker in the query. Deliberately conservative — a missed play
    session is swept by the daily runner; a false positive emails a gazette
    about governance work, which is merely embarrassing."""
    if not query:
        return False
    if swarm_gate.parse_purpose(query) == "creative-production":
        return True
    low = query.lower()
    return any(marker in low for marker in PLAY_MARKERS)


# ============================================================
# Collectors — receipts on disk, nothing else
# ============================================================

_SID_RE = re.compile(r"^(\d{8})_(\d{6})$")


def _sid_datetime(session_id: str) -> "datetime | None":
    """Session ids are timestamps (20260704_122658) — parse or None."""
    m = _SID_RE.match(session_id or "")
    if not m:
        return None
    try:
        return datetime.strptime(f"{m.group(1)}{m.group(2)}", "%Y%m%d%H%M%S")
    except ValueError:
        return None


def collect_sessions(logs_dir: Path, since: datetime,
                     until: "datetime | None" = None) -> list[dict]:
    """One record per closed session in the window, joined across the closer's
    three receipt files. The session id's own timestamp decides membership
    (deterministic; mtime lies after copies/restores)."""
    until = until or datetime.now()
    sessions: list[dict] = []
    try:
        scorecards = sorted(Path(logs_dir).glob("scorecard-*.json"))
    except OSError as e:
        logger.error(f"[gazette] cannot list {logs_dir}: {e}")
        return []
    for path in scorecards:
        sid = path.stem.removeprefix("scorecard-")
        started = _sid_datetime(sid)
        if started is None or not (since <= started < until):
            continue
        try:
            sc = json.loads(path.read_text())
        except (OSError, ValueError) as e:
            logger.warning(f"[gazette] unreadable scorecard {path.name}: {e}")
            sc = {}
        rec = {
            "session_id": sid,
            "started": started.isoformat(timespec="seconds"),
            "query": sc.get("query") or "",
            "rounds": sc.get("rounds"),
            "models": sc.get("models_active") or [],
            "persistence_gap": sc.get("persistence_gap"),
            "phantom_paths": (sc.get("filestore") or {}).get("phantom_paths") or [],
            "honest_verb_violations": (sc.get("filestore") or {}).get("honest_verb_violations"),
            "blockers": (sc.get("synthesis_directives") or {}).get("blockers"),
            "reviews": (sc.get("synthesis_directives") or {}).get("reviews"),
            "flags": (sc.get("synthesis_directives") or {}).get("flags"),
            "truncated_models": sc.get("truncated_models") or [],
            "rate_limited_models": sc.get("rate_limited_models") or [],
            "email_gap": (sc.get("audit_counts") or {}).get("gap") if sc.get("audit_counts") else None,
            "gate_failures": (sc.get("gate") or {}).get("gate_failures"),
        }
        corpus_path = Path(logs_dir) / "corpus" / f"corpus-{sid}.json"
        if corpus_path.exists():
            try:
                ev = json.loads(corpus_path.read_text())
                rec["query"] = rec["query"] or ev.get("query") or ""
                rec["repo_sha"] = ev.get("repo_sha")
            except (OSError, ValueError):
                pass
        digest_path = Path(logs_dir) / f"closer-digest-{sid}.md"
        if digest_path.exists():
            try:
                rec["needs_you"] = _needs_you_lines(digest_path.read_text())
            except OSError:
                rec["needs_you"] = []
        rec["is_play"] = is_play_session(rec["query"])
        sessions.append(rec)
    sessions.sort(key=lambda r: r["session_id"])
    return sessions


def _needs_you_lines(digest_text: str) -> list[str]:
    """The bullet lines under the digest's '## What needs you' heading —
    already human-triaged by the closer; we just carry them forward."""
    lines: list[str] = []
    in_section = False
    for line in (digest_text or "").splitlines():
        if line.startswith("## "):
            in_section = line.strip() == "## What needs you"
            continue
        if in_section and line.strip().startswith("- "):
            item = line.strip()[2:].strip()
            if item and not item.startswith("(nothing"):
                lines.append(item)
    return lines[:10]


def collect_joy_runs(since: datetime, until: "datetime | None" = None) -> list[str]:
    """Joy run folders in the window, by date prefix (joy/runs/YYYY-MM-DD*).
    Names only — the runs' own scorecards are their receipts."""
    until = until or datetime.now()
    dates = set()
    day = since.date()
    while day <= until.date():
        dates.add(day.isoformat())
        day += timedelta(days=1)
    try:
        entries = swarm_filestore.list_files("joy/runs") or []
    except Exception as e:  # list failure = report nothing, never crash the paper
        logger.warning(f"[gazette] joy/runs listing failed: {e}")
        return []
    hits = set()
    for entry in entries:
        name = str(entry)
        marker = name.split("joy/runs/", 1)[-1]
        prefix = marker.split("/", 1)[0][:10]
        if prefix in dates:
            hits.add(marker.split("/", 1)[0])
    return sorted(hits)


AUDIT_EXCERPT_CHARS = 3000


def collect_persistence_audits(since: datetime,
                               until: "datetime | None" = None) -> list[dict]:
    """Persistence audits the sessions wrote to the filestore logs/ lane
    inside the window. Returns [{path, text}] with text capped at
    AUDIT_EXCERPT_CHARS (audits are short; the cap is a torn-file guard).

    Listing is preferred (catches variant filenames); per-date direct reads
    are the fallback when listing fails — an unlistable directory must not
    silence the Audit Desk (the joy/metrics lesson)."""
    until = until or datetime.now()
    dates = []
    day = since.date()
    while day <= until.date():
        dates.append(day.isoformat())
        day += timedelta(days=1)

    paths: list[str] = []
    try:
        for entry in swarm_filestore.list_files("logs") or []:
            name = str(entry)
            if "persistence-audit" in name and any(d in name for d in dates):
                paths.append(name.split("swarm/", 1)[-1].lstrip("/"))
    except Exception as e:
        logger.warning(f"[gazette] logs listing failed ({e}); falling back to per-date reads")
    if not paths:
        paths = [f"logs/{d}_persistence-audit.md" for d in dates]

    audits: list[dict] = []
    for path in sorted(set(paths)):
        text = swarm_filestore.read_file(path)
        if text:
            audits.append({"path": path, "text": text[:AUDIT_EXCERPT_CHARS]})
    return audits


def collect_email_log(since: datetime, until: "datetime | None" = None) -> list[dict]:
    """Entries from logs/emails.log inside the window — the ground truth the
    persistence audits learned to demand. Returns [{timestamp, subject}]."""
    until = until or datetime.now()
    raw = swarm_filestore.read_file("logs/emails.log")
    if not raw:
        return []
    entries: list[dict] = []
    for block in raw.split("\n\n---\n\n"):
        lines = block.strip().splitlines()
        if not lines:
            continue
        head = lines[0].split(" | ")[0].strip()
        try:
            ts = datetime.fromisoformat(head)
        except ValueError:
            continue
        if not (since <= ts < until):
            continue
        subject = ""
        for line in lines[1:]:
            if line.startswith("subject: "):
                subject = line.removeprefix("subject: ").strip()
                break
        entries.append({"timestamp": ts.isoformat(timespec="seconds"), "subject": subject})
    return entries


# ============================================================
# Editions — pure formatters over collected receipts
# ============================================================

def _fmt(value) -> str:
    """Unmeasured renders as '?', never as 0 — same rule as the scorecard."""
    return "?" if value is None else str(value)


def _sum_measured(sessions: list[dict], key: str) -> int:
    return sum(s[key] for s in sessions if isinstance(s.get(key), int))


def build_daily_burrow(*, date_str: str, sessions: list[dict],
                       joy_runs: list[str], email_entries: list[dict],
                       audits: "list[dict] | None" = None) -> tuple[str, str]:
    """(subject, markdown body) for the daily edition. Pure."""
    n = len(sessions)
    blockers = _sum_measured(sessions, "blockers")
    flags = _sum_measured(sessions, "flags")
    email_gaps = [(s["session_id"], s["email_gap"]) for s in sessions
                  if isinstance(s.get("email_gap"), int) and s["email_gap"] > 0]
    gap_total = sum(g for _, g in email_gaps)
    play = [s for s in sessions if s.get("is_play")]

    subject = (f"Daily Burrow {date_str} — {n} session(s), {flags} flag(s), "
               f"{blockers} blocker(s)" + (f", {gap_total} unsent email(s)" if gap_total else ""))

    lines: list[str] = [
        "# The RRI Daily Burrow",
        f"Date: {date_str} · Daemon window: last 24h · gazette v{GAZETTE_VERSION}",
        "",
        "## Front Page",
    ]
    if not sessions and not joy_runs:
        lines.append("A quiet burrow: no sessions closed and no Joy runs in the window.")
    else:
        summary = [f"{n} session(s) closed"]
        if play:
            summary.append(f"{len(play)} PLAY")
        if joy_runs:
            summary.append(f"{len(joy_runs)} Joy run(s)")
        if blockers:
            summary.append(f"{blockers} blocker(s) need you")
        if gap_total:
            summary.append(f"{gap_total} owed email(s) never sent — see Email Channel")
        lines.append("; ".join(summary) + ".")

    lines += ["", "## Session Ledger",
              "| Session | Rounds | Models | Gap | Blockers | Flags | Play |",
              "|---|---:|---|---:|---:|---:|---|"]
    for s in sessions:
        lines.append(
            f"| {s['session_id']} | {_fmt(s['rounds'])} | {len(s['models'])} "
            f"| {_fmt(s['persistence_gap'])} | {_fmt(s['blockers'])} "
            f"| {_fmt(s['flags'])} | {'yes' if s['is_play'] else ''} |")
    if not sessions:
        lines.append("| (none) | | | | | | |")

    lines += ["", "## What Needs You"]
    needs = [(s["session_id"], item) for s in sessions for item in (s.get("needs_you") or [])]
    if needs:
        for sid, item in needs:
            lines.append(f"- [{sid}] {item}")
    else:
        lines.append("- (nothing flagged for your attention in this window)")

    lines += ["", "## Email Channel"]
    lines.append(f"- Logged sends in window: {len(email_entries)}")
    if email_gaps:
        for sid, gap in email_gaps:
            lines.append(f"- **Session {sid}: gap={gap} — {gap} owed email(s) did NOT send.** "
                         f"The content is in `closer-digest-{sid}.md` ('What was flagged').")
    else:
        lines.append("- No unsent-email gaps reported by closers in this window.")

    lines += ["", "## Joy / Play"]
    if joy_runs:
        for run in joy_runs:
            lines.append(f"- joy/runs/{run}")
    else:
        lines.append("- No Joy runs in the window.")
    for s in play:
        lines.append(f"- PLAY session {s['session_id']}: gazette edition in "
                     f"{PLAY_LANE}/ (attached or swept by the next Burrow).")

    lines += ["", "## Audit Desk"]
    if audits:
        for audit in audits:
            lines += [f"### {audit['path']}", "", audit["text"].strip(), ""]
    else:
        lines.append("- No persistence audits written in this window.")

    lines += ["", "## Receipts",
              "- logs/scorecard-<session>.json · logs/corpus/corpus-<session>.json · logs/closer-digest-<session>.md",
              "- logs/emails.log (send ground truth) · filestore logs/*persistence-audit*.md",
              "",
              "_Generated mechanically from closer receipts; no model narration was consulted._"]
    return subject, "\n".join(lines)


def build_play_gazette(session: dict) -> tuple[str, str]:
    """(subject, markdown body) for one PLAY session's front page. Pure."""
    sid = session["session_id"]
    clipped = bool(session.get("truncated_models"))
    status = "INCOMPLETE / CLIPPED — RESUME REQUIRED" if clipped else "complete"
    subject = f"Play Gazette — {sid}" + (" — CLIPPED" if clipped else "")

    query_head = (session.get("query") or "").strip().replace("\n", " ")[:160]
    lines = [
        "# THE RRI PLAY GAZETTE",
        f"Session {sid} · {session.get('started', '')} · gazette v{GAZETTE_VERSION}",
        "",
        "## Front Page",
        f"**{query_head or '(untitled play session)'}**",
        "",
        "## The Box Score",
        f"- Rounds: {_fmt(session.get('rounds'))}",
        f"- Models present: {', '.join(session.get('models') or []) or '?'}",
        f"- Persistence gap: {_fmt(session.get('persistence_gap'))}",
        f"- Honest-verb violations: {_fmt(session.get('honest_verb_violations'))}",
        f"- Flags/Blockers/Reviews: {_fmt(session.get('flags'))} / "
        f"{_fmt(session.get('blockers'))} / {_fmt(session.get('reviews'))}",
        "",
        "## Production Status",
        f"- Status: **{status}**",
    ]
    if clipped:
        lines.append(f"- Clipped seats: {', '.join(session['truncated_models'])} — "
                     "resume from the exact clip point in the transcript; do not re-run fresh.")
    if session.get("rate_limited_models"):
        lines.append(f"- Rate-limited seats: {', '.join(session['rate_limited_models'])}")
    if session.get("phantom_paths"):
        lines.append("- Claimed-but-unpersisted paths (do NOT cite): "
                     + ", ".join(session["phantom_paths"][:10]))

    needs = session.get("needs_you") or []
    lines += ["", "## Human Action Needed"]
    lines += [f"- {item}" for item in needs] or ["- (none)"]

    lines += ["", "## Receipts",
              f"- logs/closer-digest-{sid}.md (full digest)",
              f"- logs/scorecard-{sid}.json",
              "",
              "_Box score from closer receipts; the transcript remains the canonical text._"]
    return subject, "\n".join(lines)


# ============================================================
# DOCX rendering — optional, honest about absence
# ============================================================

def render_docx(title: str, body_md: str, out_path: Path) -> bool:
    """Render a gazette's markdown to a simple styled DOCX. Returns False
    (never raises) when python-docx is unavailable or the write fails —
    the caller then sends the email honestly unattached."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.info("[gazette] python-docx unavailable — skipping DOCX edition")
        return False
    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        head = doc.add_heading(title, level=0)
        for run in head.runs:
            run.font.color.rgb = RGBColor(0xC4, 0x65, 0x4A)

        table_rows: list[list[str]] = []

        def _flush_table():
            nonlocal table_rows
            rows = [r for r in table_rows if not set("".join(r)) <= set("-: ")]
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Light Grid Accent 2"
                for i, cells in enumerate(rows):
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell
            table_rows = []

        for line in body_md.splitlines():
            stripped = line.strip()
            if stripped.startswith("|") and stripped.endswith("|"):
                table_rows.append([c.strip() for c in stripped.strip("|").split("|")])
                continue
            _flush_table()
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith("# "):
                continue  # already the document title
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:].replace("**", ""), style="List Bullet")
            elif stripped:
                doc.add_paragraph(stripped.replace("**", "").replace("_", ""))
        _flush_table()

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return True
    except Exception as e:
        logger.error(f"[gazette] DOCX render failed for {out_path}: {type(e).__name__}: {e}")
        return False


# ============================================================
# Publish — persist the edition, then email it
# ============================================================

def _slug_date(session_id: str) -> str:
    dt = _sid_datetime(session_id)
    return dt.date().isoformat() if dt else "undated"


def play_gazette_path(session_id: str) -> str:
    return f"{PLAY_LANE}/{_slug_date(session_id)}_{session_id}-gazette.md"


def fire_play_gazette(*, session_id: str, logs_dir: Path, outputs_dir: Path) -> dict:
    """Build, persist, and email the Play Gazette for one closed session.
    Idempotent: the persisted filestore edition is the marker. Never raises."""
    result = {"session_id": session_id, "published": False, "emailed": False, "docx": None}
    try:
        marker = play_gazette_path(session_id)
        if swarm_filestore.read_file(marker) is not None:
            logger.info(f"[gazette] play edition already published for {session_id}")
            return result
        started = _sid_datetime(session_id) or datetime.now()
        records = collect_sessions(Path(logs_dir), started - timedelta(seconds=1),
                                   started + timedelta(seconds=1))
        session = next((r for r in records if r["session_id"] == session_id), None)
        if session is None:
            logger.warning(f"[gazette] no receipts found for session {session_id}; not publishing")
            return result
        subject, body = build_play_gazette(session)

        if not swarm_filestore.write_file(marker, body):
            logger.error(f"[gazette] could not persist play edition {marker}")
            return result
        result["published"] = True

        docx_path = Path(outputs_dir) / f"play_gazette_{session_id}.docx"
        attachments = [str(docx_path)] if render_docx(subject, body, docx_path) else None
        if attachments:
            result["docx"] = str(docx_path)

        import swarm_mail
        sent, reason = swarm_mail.send_operational(
            subject, body, attachments=attachments,
            prefix="[RRI Play Gazette]", session_id=session_id)
        result["emailed"] = sent
        if not sent:
            logger.warning(f"[gazette] play edition persisted but not emailed: {reason}")
    except Exception as e:
        logger.error(f"[gazette] fire_play_gazette failed for {session_id}: {type(e).__name__}: {e}")
    return result
